from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('GLS BB Platform Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph('GLS BB is a Capture The Flag (CTF) style cybersecurity training platform designed to simulate real-world vulnerability assessment scenarios. Participants operate as security researchers, identifying and exploiting vulnerabilities within a controlled environment to capture "flags" and earn points. The platform emphasizes skill development across multiple security domains, featuring real-time progress tracking and an integrated hint system.')

    # 2. Core Platform Functionalities
    doc.add_heading('2. Core Platform Functionalities', level=1)

    doc.add_heading('The Objective', level=2)
    p = doc.add_paragraph()
    p.add_run('The primary goal for participants is to locate and submit ').bold = True
    p.add_run('Flags').bold = True
    p.add_run('.')
    
    doc.add_paragraph('Definition: A flag is a specific text string acting as proof of exploitation (e.g., flag{s0m3_t3xt_h3r3}).', style='List Bullet')
    doc.add_paragraph('Mechanism: Flags are hidden within challenges, requiring participants to bypass security controls or decode information.', style='List Bullet')
    doc.add_paragraph('Scoring: Successful submission of a flag awards points to the user\'s account.', style='List Bullet')

    doc.add_heading('Leaderboard System', level=2)
    p = doc.add_paragraph('The platform maintains a real-time leaderboard to foster competition.')
    doc.add_paragraph('Ranking Criteria: Participants are ranked primarily by total points.', style='List Bullet')
    doc.add_paragraph('Tie-Breaking: In the event of equal scores, the participant who achieved the score first is ranked higher.', style='List Bullet')
    doc.add_paragraph('Visualization: The top three participants are visually distinguished to highlight leading performers.', style='List Bullet')

    doc.add_heading('Hint System', level=2)
    p = doc.add_paragraph('To assist participants who may be stuck, the platform provides a tiered hint system:')
    doc.add_paragraph('Free Hints: Basic guidance available to all participants immediately without penalty.', style='List Number')
    doc.add_paragraph('Unlockable Hints: Advanced clues that can be purchased using earned points. Using these hints introduces a strategic element, as it reduces the user\'s net score potential.', style='List Number')

    doc.add_heading('Challenge Progression', level=2)
    doc.add_paragraph('Chained Access: Certain challenges are sequential. Access to advanced challenges (e.g., Level 2) is restricted until the prerequisite (Level 1) is successfully solved. This ensures a structured learning path.')

    # 3. Challenge Domains
    doc.add_heading('3. Challenge Domains', level=1)
    doc.add_paragraph('The platform hosts challenges categorized into five distinct cybersecurity domains:')

    doc.add_heading('1. Website (Web Exploitation)', level=2)
    doc.add_paragraph('Focuses on identifying vulnerabilities in web applications.')
    doc.add_paragraph('Concept: Exploiting logic errors, injection flaws, or misconfigurations.', style='List Bullet')
    doc.add_paragraph('Example Scenario: A developer comment left in HTML source code containing sensitive credentials.', style='List Bullet')
    doc.add_paragraph('Target Skills: HTML/JavaScript analysis, SQL Injection, Cross-Site Scripting (XSS).', style='List Bullet')

    doc.add_heading('2. Cryptography', level=2)
    doc.add_paragraph('Focuses on mathematical techniques for securing information.')
    doc.add_paragraph('Concept: Decoding encrypted messages or breaking weak ciphers.', style='List Bullet')
    doc.add_paragraph('Example Scenario: A text string encoded in Base64 or shifted using a Caesar Cipher.', style='List Bullet')
    doc.add_paragraph('Target Skills: Encryption algorithms, hashing, frequency analysis.', style='List Bullet')

    doc.add_heading('3. Steganography', level=2)
    doc.add_paragraph('The practice of concealing information within non-secret text or data.')
    doc.add_paragraph('Concept: Extracting hidden data from images, audio files, or videos.', style='List Bullet')
    doc.add_paragraph('Example Scenario: A text file embedded within the binary data of a JPEG image.', style='List Bullet')
    doc.add_paragraph('Target Skills: File signature analysis, metadata inspection, hex editing.', style='List Bullet')

    doc.add_heading('4. Forensics', level=2)
    doc.add_paragraph('The application of investigation techniques to recover digital evidence.')
    doc.add_paragraph('Concept: Analyzing system logs, memory dumps, or network traffic.', style='List Bullet')
    doc.add_paragraph('Example Scenario: Reassembling a deleted file or extracting a flag from a network packet capture (.pcap).', style='List Bullet')
    doc.add_paragraph('Target Skills: Packet analysis, file recovery, log auditing.', style='List Bullet')

    doc.add_heading('5. OSINT (Open Source Intelligence)', level=2)
    doc.add_paragraph('The collection and analysis of data gathered from open sources.')
    doc.add_paragraph('Concept: Locating sensitive information exposed publicly on the internet.', style='List Bullet')
    doc.add_paragraph('Example Scenario: Finding hardcoded credentials in a public code repository or social media post.', style='List Bullet')
    doc.add_paragraph('Target Skills: Advanced search techniques, digital footprint analysis.', style='List Bullet')

    # 4. Administration Functionalities
    doc.add_heading('4. Administration Functionalities', level=1)
    doc.add_paragraph('Administrators possess comprehensive control over the event via a centralized dashboard.')

    doc.add_heading('Challenge Management', level=2)
    doc.add_paragraph('Creation: Ability to deploy new challenges with custom titles, descriptions, points, and flags.', style='List Bullet')
    doc.add_paragraph('Modification: Real-time editing of challenge details to correct errors or adjust difficulty.', style='List Bullet')
    doc.add_paragraph('Deletion: Removal of challenges from the active event.', style='List Bullet')
    doc.add_paragraph('Status Control: Admins can toggle the visibility of challenge categories (e.g., displaying "Live" vs. "Yet to Start").', style='List Bullet')

    doc.add_heading('User Management', level=2)
    doc.add_paragraph('Administrators can moderate the event to ensure integrity:')
    doc.add_paragraph('Ban/Unban: Suspend accounts suspected of violating rules or restore access.', style='List Bullet')
    doc.add_paragraph('Penalize: Deduct points from specific users for minor infractions.', style='List Bullet')
    doc.add_paragraph('Reset: Completely wipe a user\'s progress, including solves and points, effectively restarting their session.', style='List Bullet')

    doc.add_heading('Monitoring', level=2)
    doc.add_paragraph('Analytics: View detailed statistics on challenge solve rates (Locked vs. Unlocked).', style='List Bullet')
    doc.add_paragraph('Live Tracking: Monitor the leaderboard in real-time to track top performers and event progress.', style='List Bullet')

    doc.save('GLS_BB_Platform_Documentation.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
